import base64
import io

import requests
import streamlit as st
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

from charts import kpi_oee_chart_png, maturity_chart_png
from services.api_client import generate_executive_report_narrative

st.title("Executive Report")

if "workspace_id" not in st.session_state:
    st.warning("Select or create a workspace on the Home page first.")
    st.stop()

workspace_id = st.session_state["workspace_id"]

swot = st.session_state.get("swot_result")
roadmap = st.session_state.get("roadmap_result")
maturity = st.session_state.get("maturity_result")
technology = st.session_state.get("technology_result")
kpi_df = st.session_state.get("kpi_results")
kpi = kpi_df.to_dict(orient="records") if kpi_df is not None else None

st.subheader("Available findings")
status = [
    ("SWOT", swot),
    ("Digital Maturity", maturity),
    ("Roadmap", roadmap),
    ("Technology Recommendations", technology),
    ("KPI Data", kpi),
]
for label, value in status:
    st.write(f"{'✅' if value else '⬜'} {label}")

if not any(value for _, value in status):
    st.info(
        "Nothing generated yet — visit the SWOT, Roadmap, Digital Maturity, "
        "Technology, or KPI pages first, then come back here."
    )
    st.stop()

if st.button("Generate Executive Summary"):
    try:
        with st.spinner("Synthesizing findings into a narrative..."):
            narrative = generate_executive_report_narrative(
                workspace_id,
                {
                    "swot": swot,
                    "roadmap": roadmap,
                    "maturity": maturity,
                    "technology": technology,
                    "kpi": kpi,
                },
            )
    except requests.exceptions.RequestException as exc:
        st.error(f"Could not generate the report. ({exc})")
        st.stop()
    st.session_state["executive_narrative"] = narrative

if "executive_narrative" not in st.session_state:
    st.stop()

narrative = st.session_state["executive_narrative"]

MATURITY_CATEGORIES = [
    "leadership",
    "operations",
    "technology",
    "data",
    "supply_chain",
    "automation",
    "sustainability",
    "cybersecurity",
    "workforce",
]


def build_report_sections() -> list[tuple[str, str | list[str]]]:
    """One shared structure, consumed by the on-page display and all
    three export formats — so the report content is defined once."""
    sections: list[tuple[str, str | list[str]]] = [
        ("Executive Summary", narrative["executive_summary"]),
        ("Current Situation", narrative["current_situation"]),
        ("Key Findings", narrative["key_findings"]),
    ]

    if kpi:
        kpi_lines = [
            f"{row.get('machine')}: OEE {row.get('oee'):.1%}, "
            f"MTBF {row.get('mtbf_hours'):.1f}h, MTTR {row.get('mttr_hours'):.1f}h"
            for row in kpi
        ]
        sections.append(("Operational Analysis", kpi_lines))

    if swot:
        swot_lines = []
        for category in ["strengths", "weaknesses", "opportunities", "threats"]:
            swot_lines.append(f"{category.title()}:")
            for entry in swot.get(category, []):
                swot_lines.append(f"  - {entry.get('item')}: {entry.get('explanation')}")
        sections.append(("SWOT Analysis", swot_lines))

    if maturity:
        maturity_lines = [f"Overall: {maturity.get('overall')} / 5"]
        for category in MATURITY_CATEGORIES:
            entry = maturity.get(category) or {}
            label = category.replace("_", " ").title()
            maturity_lines.append(f"{label}: {entry.get('score')}/5 — {entry.get('justification')}")
        sections.append(("Digital Maturity", maturity_lines))

    if technology:
        tech_lines = [
            f"{rec.get('problem')} -> {rec.get('recommendation')} via "
            f"{rec.get('platform')} (ROI: {rec.get('expected_return')})"
            for rec in technology.get("recommendations", [])
        ]
        sections.append(("Business Opportunities (Technology Recommendations)", tech_lines))

    if roadmap:
        roadmap_lines = []
        for horizon, label in [
            ("quick_wins", "Quick Wins (0-3 months)"),
            ("medium_term", "Medium-Term (3-12 months)"),
            ("long_term", "Long-Term (12-36 months)"),
        ]:
            for item in roadmap.get(horizon, []):
                roadmap_lines.append(
                    f"[{label}] {item.get('initiative')} — "
                    f"Cost: {item.get('estimated_cost')}, ROI: {item.get('expected_return')}"
                )
        sections.append(("Transformation Roadmap", roadmap_lines))

    sections.append(("Financial Impact", narrative["financial_impact"]))
    sections.append(("Next Steps", narrative["next_steps"]))
    return sections


sections = build_report_sections()

# Charts, keyed by the section heading they belong under — generated once,
# reused across the on-page display and all three export formats, so the
# report isn't just bullet lists of numbers with no visual.
charts: dict[str, bytes] = {}
if maturity:
    charts["Digital Maturity"] = maturity_chart_png(maturity)
if kpi:
    charts["Operational Analysis"] = kpi_oee_chart_png(kpi)

for heading, content in sections:
    st.subheader(heading)
    if heading in charts:
        st.image(charts[heading])
    if isinstance(content, list):
        for line in content:
            st.write(line if line.startswith("  ") else f"- {line}")
    else:
        st.write(content)


def build_markdown(sections, charts) -> str:
    parts = ["# Executive Report\n"]
    for heading, content in sections:
        parts.append(f"## {heading}\n")
        if heading in charts:
            b64 = base64.b64encode(charts[heading]).decode("ascii")
            parts.append(f"![{heading} chart](data:image/png;base64,{b64})\n")
        if isinstance(content, list):
            for line in content:
                parts.append(f"- {line}")
        else:
            parts.append(content)
        parts.append("")
    return "\n".join(parts)


def build_docx(sections, charts) -> bytes:
    doc = Document()
    doc.add_heading("Executive Report", level=0)
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        if heading in charts:
            doc.add_picture(io.BytesIO(charts[heading]), width=Inches(6))
        if isinstance(content, list):
            for line in content:
                doc.add_paragraph(line, style="List Bullet")
        else:
            doc.add_paragraph(content)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# fpdf2's core fonts only support Latin-1 — LLM-generated prose routinely
# includes em-dashes, curly quotes, and currency symbols outside that
# range (confirmed: crashed on a euro sign during testing). Sanitize to
# safe equivalents, with errors="replace" as a final net so PDF export
# can never crash on unexpected characters.
_PDF_CHAR_REPLACEMENTS = {
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "–": "-",
    "—": "-",
    "…": "...",
    "•": "-",
    " ": " ",
}


def _sanitize_for_pdf(text: str) -> str:
    for char, replacement in _PDF_CHAR_REPLACEMENTS.items():
        text = text.replace(char, replacement)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def build_pdf(sections, charts) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Executive Report", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    for heading, content in sections:
        # multi_cell(0, ...) uses "remaining width to the right margin"
        # measured from the CURRENT x position — after a previous
        # multi_cell call, x isn't always back at the left margin, which
        # can leave the next call almost no room and crash with
        # "Not enough horizontal space to render a single character".
        # Reproduced this locally before trusting the fix.
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "B", 13)
        pdf.multi_cell(0, 8, _sanitize_for_pdf(heading))
        if heading in charts:
            pdf.set_x(pdf.l_margin)
            pdf.image(io.BytesIO(charts[heading]), x=pdf.l_margin, w=180)
            pdf.ln(2)
        pdf.set_font("Helvetica", "", 11)
        if isinstance(content, list):
            for line in content:
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, _sanitize_for_pdf(f"- {line}"))
        else:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 6, _sanitize_for_pdf(content))
        pdf.ln(3)
    return bytes(pdf.output())


st.divider()
st.subheader("Export")
col1, col2, col3 = st.columns(3)
with col1:
    st.download_button(
        "Download Markdown",
        build_markdown(sections, charts),
        file_name="executive_report.md",
        mime="text/markdown",
    )
with col2:
    st.download_button(
        "Download Word",
        build_docx(sections, charts),
        file_name="executive_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
with col3:
    st.download_button(
        "Download PDF",
        build_pdf(sections, charts),
        file_name="executive_report.pdf",
        mime="application/pdf",
    )
